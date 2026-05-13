"""
Generates professional Word (.docx) and PDF cover letters.
"""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

NAVY = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)

MONTHS_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _date_es() -> str:
    n = datetime.now()
    return f"Caracas, {n.day} de {MONTHS_ES[n.month - 1]} de {n.year}"


def _contact_line(profile: dict) -> str:
    parts = [x for x in [
        profile.get("email"), profile.get("phone"),
        profile.get("location"), profile.get("linkedin"),
    ] if x]
    return "  ·  ".join(parts)


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _p(doc: Document, text: str = "", bold: bool = False, size: int = 11,
       align=WD_ALIGN_PARAGRAPH.LEFT, color: RGBColor = None,
       space_after: int = 6) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if text:
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    return para


def to_docx(profile: dict, job: dict, letter_text: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # ── Header: Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(3)
    r = name_para.add_run(profile.get("name", "").upper())
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = NAVY

    # ── Header: Contact line ──
    contact = _contact_line(profile)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        r = cp.add_run(contact)
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY

    _hr(doc)

    # spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)

    # ── Date (right-aligned) ──
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.paragraph_format.space_after = Pt(14)
    r = dp.add_run(_date_es())
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    # ── Recipient block ──
    if job.get("company"):
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(2)
        r = cp.add_run(job["company"])
        r.font.bold = True
        r.font.size = Pt(10)

    loc = job.get("location", "")
    if loc and "remoto" not in loc.lower():
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        r = lp.add_run(loc)
        r.font.size = Pt(10)
        r.font.color.rgb = GRAY

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(8)

    # ── Subject line ──
    subj = doc.add_paragraph()
    subj.paragraph_format.space_after = Pt(14)
    r1 = subj.add_run("Ref.: ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = ACCENT
    r2 = subj.add_run(f"Postulación al cargo de {job.get('title', 'la vacante')}")
    r2.font.size = Pt(10)

    # ── Body paragraphs ──
    for chunk in letter_text.split("\n\n"):
        chunk = chunk.strip().replace("\n", " ")
        if not chunk:
            continue
        bp = doc.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.space_after = Pt(10)
        r = bp.add_run(chunk)
        r.font.size = Pt(11)

    # ── Closing ──
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _p(doc, "Atentamente,", size=11, space_after=20)

    sig = doc.add_paragraph()
    sig.paragraph_format.space_after = Pt(2)
    r = sig.add_run(profile.get("name", ""))
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    if profile.get("email"):
        _p(doc, profile["email"], size=10, color=GRAY, space_after=2)
    if profile.get("phone"):
        _p(doc, profile["phone"], size=10, color=GRAY, space_after=2)
    if profile.get("linkedin"):
        _p(doc, profile["linkedin"], size=10, color=GRAY, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _find_font(bold: bool = False) -> tuple[str, str]:
    """Return (font_name, path) for a Unicode-capable TTF on Windows."""
    regular = "C:/Windows/Fonts/calibri.ttf"
    boldf = "C:/Windows/Fonts/calibrib.ttf"
    if not os.path.exists(regular):
        regular = "C:/Windows/Fonts/arial.ttf"
        boldf = "C:/Windows/Fonts/arialbd.ttf"
    if os.path.exists(regular):
        return "CustomFont", regular, (boldf if os.path.exists(boldf) else regular)
    return "Helvetica", None, None


def to_pdf(profile: dict, job: dict, letter_text: str) -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_margins(25, 25, 25)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)

    fname, reg_path, bold_path = _find_font()
    if reg_path:
        pdf.add_font(fname, "", reg_path)
        pdf.add_font(fname, "B", bold_path)
        pdf.add_font(fname, "I", reg_path)

    def setf(bold=False, size=11, color=(0, 0, 0)):
        pdf.set_font(fname, "B" if bold else "", size)
        pdf.set_text_color(*color)

    # ── Name ──
    setf(bold=True, size=17, color=(26, 26, 46))
    pdf.cell(0, 9, profile.get("name", "").upper(), new_x="LMARGIN", new_y="NEXT", align="C")

    # ── Contact line ──
    contact = _contact_line(profile)
    if contact:
        setf(size=9, color=(102, 102, 102))
        pdf.cell(0, 5, contact, new_x="LMARGIN", new_y="NEXT", align="C")

    # ── Horizontal rule ──
    pdf.ln(2)
    pdf.set_draw_color(26, 26, 46)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(25, y, 185, y)
    pdf.ln(8)

    # ── Date ──
    setf(size=10, color=(102, 102, 102))
    pdf.cell(0, 6, _date_es(), new_x="LMARGIN", new_y="NEXT", align="R")
    pdf.ln(4)

    # ── Recipient ──
    if job.get("company"):
        setf(bold=True, size=10)
        pdf.cell(0, 6, job["company"], new_x="LMARGIN", new_y="NEXT")
    loc = job.get("location", "")
    if loc and "remoto" not in loc.lower():
        setf(size=10, color=(102, 102, 102))
        pdf.cell(0, 5, loc, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # ── Subject ──
    setf(bold=True, size=10, color=(192, 57, 43))
    pdf.cell(18, 6, "Ref.:")
    setf(size=10)
    pdf.cell(0, 6, f" Postulacion al cargo de {job.get('title', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(7)

    # ── Body ──
    setf(size=11)
    for chunk in letter_text.split("\n\n"):
        chunk = chunk.strip().replace("\n", " ")
        if chunk:
            pdf.multi_cell(0, 6, chunk, align="J")
            pdf.ln(3)

    # ── Closing ──
    pdf.ln(5)
    setf(size=11)
    pdf.cell(0, 6, "Atentamente,", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(12)

    setf(bold=True, size=11, color=(26, 26, 46))
    pdf.cell(0, 6, profile.get("name", ""), new_x="LMARGIN", new_y="NEXT")

    setf(size=10, color=(102, 102, 102))
    for field in ["email", "phone", "linkedin"]:
        val = profile.get(field, "")
        if val:
            pdf.cell(0, 5, val, new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
